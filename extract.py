import subprocess
import os
from zipfile import ZipFile
import shutil
# used python-docx 0.8
from docx import Document

document = Document()
# объявление переменных

package_temp_name = 'temp'
package_res_name = 'test'
package_error_name = '!ERROR'
password_zip = 'Moreum saver'
password_zip = bytes(password_zip, 'utf-8')
package_output_name = '!RESULT'

# результаты будут в папке result
#Очищаем предыдущие результаты, чтоб не мешались
if os.path.exists(package_output_name):
    shutil.rmtree(package_output_name)
#Нужен в случае первого запуска программы
if not os.path.exists(package_output_name):
    os.mkdir(package_output_name)
if not os.path.exists(package_error_name):
    os.mkdir(package_error_name)
# для всех файлов в папке проверяется на то, что они di3 и разархивируется, иначе переходит к следующему файлу
# файлы берутся из папки data
for package in os.listdir(os.getcwd() + '/data'):
    if '.di3' in package:
        #	try:
        package_name = 'data/' + package
        zip_file = ZipFile(package_name)
        zip_file.setpassword(password_zip)
        package_res_name = '!RESULT/' + package
        # создание временной папки и разархивирование туда содержимого архива
        if not os.path.exists(package_res_name):
            os.mkdir(package_res_name)
        for item in zip_file.infolist():
            filepath = item.filename
            filename = os.path.basename(filepath)
            try:
                unzip_test = zip_file.read(filepath, pwd=password_zip)
            except :
                print("ERROR ON " + package)
                shutil.move("data/" + package, package_error_name + "/" + package)
                break

            unpack_test = open(package_res_name + "/" + filename, 'wb+')
            unpack_test.write(unzip_test)
            unpack_test.close()

        if (os.path.exists(package_res_name + '/que1') and
                os.path.exists(package_res_name + '/que2') and
                os.path.exists(package_res_name + '/ans1') and
                os.path.exists(package_res_name + '/ans2') and
                os.path.exists(package_res_name + '/res')):
            # Считываем вопросы и ответы из каждого файла
            f = open(package_res_name + '/que1', 'r', encoding='cp1251')
            document.add_heading('Вопрос 1', level=1)
            for line in f:
                document.add_paragraph(line)
            f.close()

            f = open(package_res_name + '/ans1', 'r', encoding='cp1251')
            document.add_heading('Ответ:', level=1)
            for line in f:
                document.add_paragraph(line)
            f.close()

            f = open(package_res_name + '/que2', 'r', encoding='cp1251')
            document.add_heading('Вопрос 2', level=1)
            for line in f:
                document.add_paragraph(line)
            f.close()

            f = open(package_res_name + '/ans2', 'r', encoding='cp1251')
            document.add_heading('Ответ:', level=1)
            for line in f:
                document.add_paragraph(line)
            f.close()

            document.save(package_res_name + '.docx')
            document = Document()
        else:
            if os.path.exists(package_res_name + '/res'):
                f = open(package_res_name + '/res', 'r', encoding='cp1251')
                i = 0
                for line in f:
                    if i==8:
                        document.save(package_output_name+"/ТЕСТ НА " + line + " БАЛЛОВ, "+package+'.docx')
                        document=Document()
                    i = i + 1

        # удаление временной папки
        if os.path.exists(package_res_name):
            path = os.path.join(os.path.abspath(os.path.dirname(__file__)), package_res_name)
            shutil.rmtree(path)

# удаление мусора с предыдущей работы программы
if os.path.exists(package_temp_name):
    path = os.path.join(os.path.abspath(os.path.dirname(__file__)), package_temp_name)
    shutil.rmtree(path)
